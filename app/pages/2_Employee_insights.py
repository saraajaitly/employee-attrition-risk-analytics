import streamlit as st
import pandas as pd
from io import BytesIO
import shap
from docx import Document
from docx.shared import Pt
from predictor import load_model, preprocess_data
def explain_feature(feature, employee):

    if feature == "OverTime_Yes":
        return "Employee frequently works overtime."

    elif feature == "OverTime_No":
        return "Employee does not regularly work overtime."

    elif feature == "WorkLifeBalance":
        levels = {
            1: "poor",
            2: "fair",
            3: "good",
            4: "excellent"
        }
        level = levels.get(employee["WorkLifeBalance"], employee["WorkLifeBalance"])
        return f"Work-life balance is {level}."

    elif feature == "EnvironmentSatisfaction":
        levels = {
            1: "low",
            2: "moderate",
            3: "high",
            4: "very high"
        }
        level = levels.get(employee["EnvironmentSatisfaction"], employee["EnvironmentSatisfaction"])
        return f"Work environment satisfaction is {level}."

    elif feature == "JobSatisfaction":
        levels = {
            1: "low",
            2: "moderate",
            3: "high",
            4: "very high"
        }
        level = levels.get(employee["JobSatisfaction"], employee["JobSatisfaction"])
        return f"Job satisfaction is {level}."

    elif feature == "RelationshipSatisfaction":
        levels = {
            1: "low",
            2: "moderate",
            3: "high",
            4: "very high"
        }
        level = levels.get(employee["RelationshipSatisfaction"], employee["RelationshipSatisfaction"])
        return f"Relationship satisfaction is {level}."

    elif feature == "JobInvolvement":
        levels = {
            1: "low",
            2: "moderate",
            3: "high",
            4: "very high"
        }
        level = levels.get(employee["JobInvolvement"], employee["JobInvolvement"])
        return f"Job involvement is {level}."

    elif feature == "YearsSinceLastPromotion":
        years = employee["YearsSinceLastPromotion"]

        if years == 0:
            return "Employee was promoted recently."

        return f"Employee has not received a promotion for {years} years."

    elif feature == "YearsAtCompany":
        years = employee["YearsAtCompany"]
        return f"Employee has been with the company for {years} years."

    elif feature == "YearsWithCurrManager":
        years = employee["YearsWithCurrManager"]
        return f"Employee has reported to the current manager for {years} years."

    elif feature == "TotalWorkingYears":
        years = employee["TotalWorkingYears"]
        return f"Employee has {years} years of total work experience."

    elif feature == "NumCompaniesWorked":
        n = employee["NumCompaniesWorked"]

        if n == 0:
            return "This is the employee's first company."
        elif n == 1:
            return "Employee has worked at one previous company."
        else:
            return f"Employee has worked at {n} previous companies."

    elif feature == "MonthlyIncome":
        return f"Monthly income is ₹{employee['MonthlyIncome']:,.0f}."

    elif feature == "StockOptionLevel":
        level = employee["StockOptionLevel"]

        if level == 0:
            return "Employee has no stock options."
        elif level == 1:
            return "Employee has basic stock option benefits."
        elif level == 2:
            return "Employee has moderate stock option benefits."
        else:
            return "Employee has extensive stock option benefits."

    elif feature == "DistanceFromHome":
        distance = employee["DistanceFromHome"]
        return f"Employee lives {distance} km from the workplace."

    elif feature == "Age":
        return f"Employee is {employee['Age']} years old."

    elif feature == "BusinessTravel_Travel_Frequently":
        return "Employee frequently travels for work."

    elif feature == "BusinessTravel_Travel_Rarely":
        return "Employee occasionally travels for work."

    elif feature == "BusinessTravel_Non-Travel":
        return "Employee rarely travels for work."

    else:
        return feature.replace("_", " ")

def recommend_action(feature, employee):

    if feature == "OverTime_Yes":
        return "Review workload and reduce overtime where possible."

    elif feature == "WorkLifeBalance":

        if employee["WorkLifeBalance"] <= 2:
            return "Discuss flexible work arrangements and work-life balance improvements."

        return "Continue supporting healthy work-life balance practices."

    elif feature == "EnvironmentSatisfaction":

        if employee["EnvironmentSatisfaction"] <= 2:
            return "Identify workplace concerns and improve the employee's work environment."

        return "Maintain the positive work environment."

    elif feature == "JobSatisfaction":

        if employee["JobSatisfaction"] <= 2:
            return "Schedule a one-on-one meeting to understand employee concerns."

        return "Continue recognizing employee contributions."

    elif feature == "RelationshipSatisfaction":

        if employee["RelationshipSatisfaction"] <= 2:
            return "Encourage regular manager check-ins and team engagement."

        return "Maintain healthy workplace relationships."

    elif feature == "JobInvolvement":

        if employee["JobInvolvement"] <= 2:
            return "Provide meaningful projects and increase employee engagement."

        return "Continue supporting employee involvement."

    elif feature == "YearsSinceLastPromotion":

        if employee["YearsSinceLastPromotion"] >= 5:
            return "Discuss career progression or promotion opportunities."

        return "Continue monitoring career growth."

    elif feature == "StockOptionLevel":

        if employee["StockOptionLevel"] == 0:
            return "Review employee benefits and long-term retention incentives."

        return "Continue offering competitive benefits."

    elif feature == "MonthlyIncome":
        return "Review compensation during the next performance cycle."

    elif feature == "DistanceFromHome":

        if employee["DistanceFromHome"] >= 20:
            return "Consider flexible or hybrid work options where feasible."

        return "No immediate action required."

    elif feature == "BusinessTravel_Travel_Frequently":
        return "Monitor travel frequency and employee well-being."

    elif feature == "YearsWithCurrManager":

        if employee["YearsWithCurrManager"] >= 8:
            return "Review career development and growth opportunities."

        return "Continue regular manager feedback sessions."

    elif feature == "YearsAtCompany":

        if employee["YearsAtCompany"] >= 10:
            return "Discuss long-term career planning and retention."

        return "Continue supporting career development."

    elif feature == "NumCompaniesWorked":

        if employee["NumCompaniesWorked"] >= 5:
            return "Conduct periodic engagement check-ins."

        return "Continue regular employee engagement."

    else:
        return None

st.set_page_config(page_title="Employee Insights", page_icon="👤", layout="wide")

st.title("👤 Employee Insights")
st.markdown("Analyze an individual employee's attrition risk and receive personalized HR recommendations.")
results = st.session_state.get("results")

if results is None:
    st.warning("⚠️ Please upload a dataset and run predictions first.")
    st.stop()
st.success("✅ Prediction results loaded successfully!")
st.subheader("Select Employee")

employee_options = (
    results["EmployeeNumber"].astype(str)
    + " - "
    + results["JobRole"]
)

selected_employee = st.selectbox(
    "Choose an Employee",
    employee_options
)
employee_id = int(selected_employee.split(" - ")[0])

employee = results[results["EmployeeNumber"] == employee_id].iloc[0]


st.subheader("👤 Employee Profile")

col1, col2, col3 = st.columns([1.95, 1, 1])

with col1:
    st.metric("Employee ID", employee["EmployeeNumber"])
    st.metric("Department", employee["Department"])
    st.metric("Job Role", employee["JobRole"])

with col2:
    st.metric("Age", employee["Age"])
    st.metric("Years at Company", employee["YearsAtCompany"])
    st.metric("Monthly Income", f"₹ {employee['MonthlyIncome']:,}")

with col3:
    st.metric("Overtime", employee["OverTime"])
    st.metric("Work-Life Balance", employee["WorkLifeBalance"])
    st.metric("Job Satisfaction", employee["JobSatisfaction"])

st.divider()

st.subheader("🎯 Attrition Prediction")
col1, col2 = st.columns(2)
with col1:
    st.metric(
        "Attrition Risk",
        f"{employee['Attrition Risk (%)']:.2f}%"
    )

with col2:
    risk = employee["Risk Category"]

    if risk == "High":
        st.error("🔴 High Risk")

    elif risk == "Medium":
        st.warning("🟡 Medium Risk")

    else:
        st.success("🟢 Low Risk")

st.divider()
model = load_model()

employee_df = employee.to_frame().T

processed_employee = preprocess_data(employee_df)

explainer = shap.TreeExplainer(model)

shap_values = explainer.shap_values(processed_employee)
feature_names = processed_employee.columns

shap_df = pd.DataFrame({
    "Feature": feature_names,
    "SHAP Value": shap_values[0]
})

shap_df["Importance"] = shap_df["SHAP Value"].abs()

risk = employee["RiskCategory"] if "RiskCategory" in employee else employee["Risk Category"]



feature_labels = {
    "OverTime_Yes": "Frequent overtime",
    "WorkLifeBalance": "Poor work-life balance",
    "JobSatisfaction": "Low job satisfaction",
    "EnvironmentSatisfaction": "Low work environment satisfaction",
    "YearsSinceLastPromotion": "Long time since last promotion",
    "MonthlyIncome": "Monthly income",
    "StockOptionLevel": "Stock option level",
    "DistanceFromHome": "Long commute distance",
    "BusinessTravel_Travel_Frequently": "Frequent business travel",
    "NumCompaniesWorked": "Multiple previous employers",
    "TotalWorkingYears": "Total working years",
    "Department_Research & Development": "Works in the Research & Development department",
    "Department_Sales": "Works in the Sales department",
}
risk = employee["Risk Category"]
if risk == "Low":
    st.subheader("🔍 Why is this employee likely to stay?")
else:
    st.subheader("🔍 Why is this employee at risk?")

if risk == "Low":
    top_features = shap_df.sort_values("SHAP Value")
else:
    top_features = shap_df.sort_values("SHAP Value", ascending=False)

if risk == "Low":
    st.success("Key factors supporting employee retention")
else:
    st.warning("Key factors increasing attrition risk")
IGNORE_FEATURES = {
    "Department_Research & Development",
    "Department_Sales",
    "Department_Human Resources",

    "Education",
    "EducationField_Life Sciences",
    "EducationField_Marketing",
    "EducationField_Medical",
    "EducationField_Other",
    "EducationField_Technical Degree",
    "EducationField_Human Resources",

    "MaritalStatus_Married",
    "MaritalStatus_Single",
    "MaritalStatus_Divorced",

    "Gender_Male",

    "EmployeeNumber",
    "EmployeeCount",
    "StandardHours",
    "Over18_Y",
}
displayed = 0

for feature in top_features["Feature"]:

    if feature in IGNORE_FEATURES:
        continue

    explanation = explain_feature(feature, employee)
    st.write(f"• {explanation}")

    displayed += 1

    if displayed == 3:
        break

st.divider()
st.subheader("💡 Recommended Actions")
displayed = 0
shown = set()

for feature in top_features["Feature"]:

    if (
        feature in IGNORE_FEATURES
        or feature.startswith("Department_")
        or feature.startswith("EducationField_")
        or feature.startswith("MaritalStatus_")
        or feature.startswith("Gender_")
    ):
        continue

    recommendation = recommend_action(feature, employee)

    if recommendation is None:
        continue

    if recommendation in shown:
        continue

    st.write(f"• {recommendation}")

    shown.add(recommendation)
    displayed += 1

    if displayed == 3:
        break

st.divider()
doc = Document()

title = doc.add_heading("Employee Attrition Report", level=1)
title.style.font.name = "Calibri"
title.style.font.size = Pt(18)

doc.add_heading("Employee Details", level=2)

doc.add_paragraph(f"Employee ID: {employee['EmployeeNumber']}")
doc.add_paragraph(f"Department: {employee['Department']}")
doc.add_paragraph(f"Job Role: {employee['JobRole']}")
doc.add_paragraph(f"Age: {employee['Age']}")
doc.add_paragraph(f"Years at Company: {employee['YearsAtCompany']}")

doc.add_heading("Attrition Prediction", level=2)

doc.add_paragraph(f"Risk Score: {employee['Attrition Risk (%)']:.2f}%")
doc.add_paragraph(f"Risk Category: {employee['Risk Category']}")

doc.add_heading("Key Factors", level=2)
displayed = 0

for feature in top_features["Feature"]:

    if (
        feature in IGNORE_FEATURES
        or feature.startswith("Department_")
        or feature.startswith("EducationField_")
        or feature.startswith("MaritalStatus_")
        or feature.startswith("Gender_")
    ):
        continue

    explanation = explain_feature(feature, employee)

    doc.add_paragraph(explanation, style="List Bullet")

    displayed += 1

    if displayed == 3:
        break


doc.add_heading("Recommended Actions", level=2)

displayed = 0
shown = set()

for feature in top_features["Feature"]:

    if (
        feature in IGNORE_FEATURES
        or feature.startswith("Department_")
        or feature.startswith("EducationField_")
        or feature.startswith("MaritalStatus_")
        or feature.startswith("Gender_")
    ):
        continue

    recommendation = recommend_action(feature, employee)

    if recommendation is None:
        continue

    if recommendation in shown:
        continue

    doc.add_paragraph(recommendation, style="List Bullet")

    shown.add(recommendation)
    displayed += 1

    if displayed == 3:
        break

buffer = BytesIO()
doc.save(buffer)
buffer.seek(0)

st.download_button(
    label="📄 Download Employee Report",
    data=buffer,
    file_name=f"employee_{employee['EmployeeNumber']}_report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True
)